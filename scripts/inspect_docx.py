"""Скрипт для анализа структуры docx файла"""
from docx import Document
import sys

def inspect_docx(filepath: str):
    doc = Document(filepath)

    print(f"=== Анализ файла: {filepath} ===\n")

    # Параграфы
    print(f"Параграфов: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  [{i}] {para.text[:80]}...")

    # Таблицы
    print(f"\nТаблиц: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- Таблица {t_idx} ({len(table.rows)} строк x {len(table.columns)} столбцов) ---")
        for r_idx, row in enumerate(table.rows):
            cells_text = [cell.text.strip()[:30] for cell in row.cells]
            print(f"  Строка {r_idx}: {cells_text}")

if __name__ == "__main__":
    filepath = sys.argv[1] if len(sys.argv) > 1 else "Анкета участника закупки.docx"
    inspect_docx(filepath)
