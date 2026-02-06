"""
Скрипт для заполнения реквизитов в docx документах.

Поддерживает два режима:
1. Замена плейсхолдеров типа {{ключ}} на значения из JSON
2. Заполнение таблиц где первый столбец - название поля
"""
import json
import re
import sys
from pathlib import Path
from docx import Document


def load_requisites(json_path: str) -> dict:
    """Загрузка реквизитов из JSON файла"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def replace_placeholders(doc: Document, data: dict) -> int:
    """
    Замена плейсхолдеров {{key}} на значения из словаря.
    Возвращает количество замен.
    """
    count = 0

    # Замена в параграфах
    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))
                count += 1

    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                        count += 1

    return count


def fill_table_by_labels(doc: Document, data: dict, label_mapping: dict,
                         label_col: int = 1, value_col: int = 2,
                         skip_headers: bool = True) -> int:
    """
    Заполнение таблиц по меткам.
    label_mapping: {"Текст метки в документе": "ключ_в_json"}
    label_col: индекс столбца с метками (по умолчанию 1 - второй столбец)
    value_col: индекс столбца для данных (по умолчанию 2 - третий столбец)
    skip_headers: пропускать строки-заголовки
    Возвращает количество заполненных ячеек.
    """
    count = 0
    # Слова которые указывают на заголовок таблицы
    header_markers = ["№", "п/п", "наименование", "сведения"]

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) > max(label_col, value_col):
                # Пропускаем первую строку (заголовок)
                if skip_headers and row_idx == 0:
                    continue

                label_text = row.cells[label_col].text.strip().lower()

                # Пропускаем если это похоже на заголовок
                if skip_headers:
                    first_cell = row.cells[0].text.strip().lower()
                    if any(m in first_cell for m in header_markers):
                        continue

                # Ищем соответствие метки (точное или частичное)
                matched = False
                for label, json_key in label_mapping.items():
                    if label.lower() in label_text:
                        if json_key in data:
                            row.cells[value_col].text = str(data[json_key])
                            count += 1
                            matched = True
                            break

                if matched:
                    continue

    return count


def fill_document(template_path: str, output_path: str, requisites_path: str, label_mapping: dict = None):
    """Основная функция заполнения документа"""

    # Загружаем данные
    data = load_requisites(requisites_path)
    print(f"Загружено {len(data)} реквизитов из {requisites_path}")

    # Открываем документ
    doc = Document(template_path)
    print(f"Открыт шаблон: {template_path}")

    # Замена плейсхолдеров
    placeholder_count = replace_placeholders(doc, data)
    print(f"Заменено плейсхолдеров: {placeholder_count}")

    # Заполнение по меткам таблиц
    if label_mapping:
        label_count = fill_table_by_labels(doc, data, label_mapping)
        print(f"Заполнено по меткам: {label_count}")

    # Сохраняем результат
    doc.save(output_path)
    print(f"Сохранено: {output_path}")


# Маппинг меток таблицы на ключи JSON
# Порядок важен! Более специфичные паттерны должны быть первыми.
# Настройте под ваш документ
DEFAULT_LABEL_MAPPING = {
    # Наименование компании (более специфичные сначала)
    "фирменное наименование": "company_name",
    "наименование, фирменное": "company_name",

    # Адреса
    "сведения о месте нахождения": "address",
    "сведения о почтовом адресе": "address",
    "место нахождения": "address",
    "почтовый адрес": "address",

    # ИНН/КПП
    "инн и кпп налогоплательщика": "inn",
    "инн и кпп": "inn",

    # Банковские реквизиты
    "банковские реквизиты": "bank_details",

    # Контакты
    "телефон участника": "phone",
    "адрес электронной почты": "email",

    # Руководитель/Ответственный
    "фамилия, имя и отчество ответст": "contact_person",
    "фамилия, имя, отчество": "director",
}


if __name__ == "__main__":
    # Пример использования
    template = sys.argv[1] if len(sys.argv) > 1 else "Анкета участника закупки.docx"
    output = sys.argv[2] if len(sys.argv) > 2 else "filled_output.docx"
    requisites = sys.argv[3] if len(sys.argv) > 3 else "requisites.json"

    fill_document(
        template_path=template,
        output_path=output,
        requisites_path=requisites,
        label_mapping=DEFAULT_LABEL_MAPPING
    )
