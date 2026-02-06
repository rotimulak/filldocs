"""Сервис заполнения docx документов"""
from pathlib import Path
from docx import Document
from typing import Optional, Union
import re


# Составные поля (несколько ключей в одной метке)
# Проверяются ПЕРВЫМИ (более специфичные)
COMPOSITE_FIELDS = {
    "инн и кпп": ["inn", "kpp"],
    "инн / кпп": ["inn", "kpp"],
    "инн/кпп": ["inn", "kpp"],
    "инн, кпп": ["inn", "kpp"],
    "р/с, к/с": ["account", "corr_account"],
    "р/с и к/с": ["account", "corr_account"],
    "серия и номер": ["passport_series", "passport_number"],
    "серия, номер": ["passport_series", "passport_number"],
    "серия/номер": ["passport_series", "passport_number"],
}

# Маппинг меток таблицы на ключи JSON
# Порядок важен! Более специфичные паттерны должны быть первыми.
DEFAULT_LABEL_MAPPING = {
    # Наименование компании
    "фирменное наименование": "company_name",
    "наименование, фирменное": "company_name",
    "наименование участника": "company_name",
    "наименование организации": "company_name",
    "полное наименование": "company_name",

    # Адреса
    "сведения о месте нахождения": "address",
    "сведения о почтовом адресе": "postal_address",
    "место нахождения": "address",
    "юридический адрес": "address",
    "почтовый адрес": "postal_address",
    "фактический адрес": "postal_address",

    # ИНН/КПП отдельно
    "инн налогоплательщика": "inn",
    "инн": "inn",
    "кпп": "kpp",

    # ОГРН
    "огрн": "ogrn",
    "огрнип": "ogrn",

    # Банковские реквизиты
    "банковские реквизиты": "bank_details",
    "реквизиты банка": "bank_details",
    "наименование банка": "bank_name",
    "название банка": "bank_name",
    "банк": "bank_name",
    "бик": "bik",
    "расчетный счет": "account",
    "расчётный счёт": "account",
    "р/с": "account",
    "корреспондентский счет": "corr_account",
    "корреспондентский счёт": "corr_account",
    "корр. счет": "corr_account",
    "корр. счёт": "corr_account",
    "к/с": "corr_account",

    # Контакты
    "телефон участника": "phone",
    "контактный телефон": "phone",
    "телефон": "phone",
    "адрес электронной почты": "email",
    "электронная почта": "email",
    "e-mail": "email",
    "email": "email",

    # Руководитель/Ответственный
    "фамилия, имя и отчество ответст": "contact_person",
    "контактное лицо": "contact_person",
    "конт. лицо": "contact_person",
    "ответственное лицо": "contact_person",
    "фамилия, имя, отчество": "director",
    "руководитель": "director",
    "директор": "director",
    "генеральный директор": "director",

    # Учредитель
    "учредитель": "founder",
    "учредители": "founder",

    # Свидетельство о регистрации
    "свидетельство о внесении в единый государственный реестр": "registration_certificate",
    "свидетельство о регистрации": "registration_certificate",
    "свидетельство егрюл": "registration_certificate",
    "свидетельство егрип": "registration_certificate",

    # Паспортные данные (полные)
    "паспортные данные": "passport",
    "паспорт": "passport",
    "данные паспорта": "passport",
    "данные документа, удостоверяющего личность": "passport",

    # Паспорт - отдельные поля
    "серия паспорта": "passport_series",
    "серия": "passport_series",
    "номер паспорта": "passport_number",

    "дата рождения": "birth_date",
    "место рождения": "birth_place",

    "кем выдан": "passport_issued_by",
    "выдан": "passport_issued_by",
    "кем выдан паспорт": "passport_issued_by",

    "код подразделения": "passport_department_code",
    "код подр.": "passport_department_code",

    "дата выдачи": "passport_issue_date",
    "дата выдачи паспорта": "passport_issue_date",
    "когда выдан": "passport_issue_date",
}

# Валидаторы полей (мягкая валидация)
VALIDATORS = {
    "inn": lambda v: len(re.sub(r'\D', '', v)) in [10, 12],
    "kpp": lambda v: len(re.sub(r'\D', '', v)) == 9,
    "ogrn": lambda v: len(re.sub(r'\D', '', v)) in [13, 15],
    "bik": lambda v: len(re.sub(r'\D', '', v)) == 9,
    "account": lambda v: len(re.sub(r'\D', '', v)) == 20,
    "corr_account": lambda v: len(re.sub(r'\D', '', v)) == 20,
    "email": lambda v: "@" in v and "." in v.split("@")[-1],
    "phone": lambda v: bool(re.match(r'^[\+8][\d\s\-\(\)]+$', v.strip())),
}

# Экстракторы - извлекают значение из "грязного" текста ячейки
# Применяются ДО нормализации и валидации
def _extract_numeric_field(text: str, valid_lengths: list[int]) -> str:
    """Извлечь первую группу цифр подходящей длины"""
    # Ищем все группы подряд идущих цифр (возможно с пробелами внутри)
    # Сначала пробуем найти группы без пробелов
    groups = re.findall(r'\d+', text)

    # Проверяем каждую группу на соответствие допустимой длине
    for group in groups:
        if len(group) in valid_lengths:
            return group

    # Если не нашли - пробуем объединить соседние группы (для случая "631 807 3604")
    all_digits = re.sub(r'\D', '', text)
    for length in sorted(valid_lengths, reverse=True):  # Сначала пробуем более длинные
        if len(all_digits) >= length:
            candidate = all_digits[:length]
            return candidate

    # Возвращаем как есть, валидация покажет ошибку
    return text


EXTRACTORS = {
    # ИНН: 10 цифр (юрлицо) или 12 цифр (ИП)
    "inn": lambda v: _extract_numeric_field(v, [10, 12]),
    # КПП: ровно 9 цифр
    "kpp": lambda v: _extract_numeric_field(v, [9]),
    # ОГРН: 13 цифр (юрлицо) или 15 цифр (ИП)
    "ogrn": lambda v: _extract_numeric_field(v, [13, 15]),
    # БИК: ровно 9 цифр
    "bik": lambda v: _extract_numeric_field(v, [9]),
    # Счета: ровно 20 цифр
    "account": lambda v: _extract_numeric_field(v, [20]),
    "corr_account": lambda v: _extract_numeric_field(v, [20]),
}

# Нормализация полей (применяется ПОСЛЕ экстракции)
NORMALIZERS = {
    "inn": lambda v: re.sub(r'\D', '', v),
    "kpp": lambda v: re.sub(r'\D', '', v),
    "ogrn": lambda v: re.sub(r'\D', '', v),
    "bik": lambda v: re.sub(r'\D', '', v),
    "account": lambda v: re.sub(r'[\s\-]', '', v),
    "corr_account": lambda v: re.sub(r'[\s\-]', '', v),
    "email": lambda v: v.strip().lower(),
}

# Паттерны для извлечения из свободного текста (fallback)
# Порядок важен - парсинг идёт с начала ячейки, первые найденные значения приоритетнее
FREE_TEXT_PATTERNS = {
    "account": r"[Рр]асч[её]тн\w*\s*сч[её]т[:\s]*(\d{20})",
    "bik": r"БИК[:\s]*(\d{9})",
    "corr_account": r"[Кк]орр(?:еспондент)?\w*\s*сч[её]т[:\s]*(\d{20})",
    "inn": r"ИНН[:\s]*(\d{10}|\d{12})",
    "kpp": r"КПП[:\s]*(\d{9})",
    "ogrn": r"ОГРН(?:ИП)?[:\s]*(\d{13}|\d{15})",
    "phone": r"[Тт]елефон[:\s]*([\d\s\-\(\)\+]{10,20})",
}

# Минимальное количество совпадений для идентификации блока реквизитов
MIN_KEYWORD_MATCHES = 3


class DocxFiller:
    """Класс для заполнения docx документов"""

    def __init__(self, label_mapping: dict = None, composite_fields: dict = None):
        self.label_mapping = label_mapping or DEFAULT_LABEL_MAPPING
        self.composite_fields = composite_fields or COMPOSITE_FIELDS

    def _count_keyword_matches(self, table) -> int:
        """Подсчёт совпадений ключевых слов в таблице"""
        count = 0
        all_patterns = list(self.composite_fields.keys()) + list(self.label_mapping.keys())

        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip().lower()
                for pattern in all_patterns:
                    if pattern in text:
                        count += 1
                        break  # Одно совпадение на ячейку
        return count

    def _find_requisites_table(self, doc: Document):
        """Найти таблицу с реквизитами (с наибольшим числом совпадений)"""
        if not doc.tables:
            return None, 0

        best_table = None
        best_score = 0

        for table in doc.tables:
            score = self._count_keyword_matches(table)
            if score > best_score:
                best_score = score
                best_table = table

        return best_table, best_score

    def _get_label_and_value_columns(self, table) -> tuple[int, int]:
        """Определить индексы колонок метки и значения"""
        if not table.rows:
            return (0, 1)

        num_cols = len(table.rows[0].cells)

        if num_cols == 2:
            return (0, 1)
        elif num_cols >= 3:
            # Проверяем, является ли первая колонка номерами (№ п/п)
            first_col_is_numbers = True
            for row in table.rows[1:]:  # Пропуск заголовка
                if row.cells:
                    text = row.cells[0].text.strip()
                    if text and not text.isdigit() and text not in ["№", ""]:
                        first_col_is_numbers = False
                        break

            if first_col_is_numbers:
                return (1, 2)
            else:
                return (0, 1)
        else:
            return (0, min(1, num_cols - 1))

    def _find_matching_keys(self, label: str) -> list[str]:
        """
        Найти ключи JSON для метки.
        Для составных полей возвращает несколько ключей.
        """
        label_lower = label.lower()

        # Сначала проверяем составные поля (более специфичные)
        for pattern, keys in self.composite_fields.items():
            if pattern in label_lower:
                return keys

        # Затем одиночные поля
        for pattern, key in self.label_mapping.items():
            if pattern in label_lower:
                return [key]

        return []

    def _extract_value(self, key: str, value: str) -> str:
        """Извлечение значения из 'грязного' текста ячейки"""
        if key in EXTRACTORS:
            return EXTRACTORS[key](value)
        return value

    def _normalize_value(self, key: str, value: str) -> str:
        """Нормализация значения поля"""
        if key in NORMALIZERS:
            return NORMALIZERS[key](value)
        return value.strip()

    def _validate_value(self, key: str, value: str) -> Optional[str]:
        """Валидация значения. Возвращает предупреждение или None"""
        if key in VALIDATORS and value:
            if not VALIDATORS[key](value):
                return f"Некорректный формат {key}: {value}"
        return None

    def _extract_from_free_text(self, text: str, used_keys: set) -> dict:
        """
        Извлечение реквизитов из свободного текста (fallback).
        Парсит с начала текста - первые найденные значения приоритетнее.
        Возвращает {"requisites": {...}, "warnings": [...]}
        """
        requisites = {}
        warnings = []

        for key, pattern in FREE_TEXT_PATTERNS.items():
            if key in used_keys:
                continue

            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                raw_value = match.group(1).strip()

                # Применяем экстракцию и нормализацию
                extracted = self._extract_value(key, raw_value)
                normalized = self._normalize_value(key, extracted)

                requisites[key] = normalized

                # Валидация
                warning = self._validate_value(key, normalized)
                if warning:
                    warnings.append(warning)

        return {"requisites": requisites, "warnings": warnings}

    def extract(self, doc_path: Path) -> dict:
        """
        Извлечение реквизитов из документа.

        Возвращает:
        {
            "requisites": {...},
            "raw_fields": [...],
            "warnings": [...],
            "table_found": bool
        }
        """
        doc = Document(str(doc_path))

        # Поиск таблицы с реквизитами
        table, score = self._find_requisites_table(doc)

        if not table or score < MIN_KEYWORD_MATCHES:
            return {
                "requisites": {},
                "raw_fields": [],
                "warnings": [f"Блок реквизитов не найден (совпадений: {score}, нужно: {MIN_KEYWORD_MATCHES})"],
                "table_found": False
            }

        label_col, value_col = self._get_label_and_value_columns(table)

        requisites = {}
        raw_fields = []
        warnings = []
        used_keys = set()

        for row in table.rows:
            if len(row.cells) <= max(label_col, value_col):
                continue

            label_text = row.cells[label_col].text.strip()
            value_text = row.cells[value_col].text.strip()

            if not label_text:
                continue

            # Поиск ключей для метки
            keys = self._find_matching_keys(label_text)

            if not keys:
                # Метка не распознана
                if value_text:
                    raw_fields.append({
                        "label": label_text,
                        "value": value_text,
                        "matched_key": None
                    })
                continue

            # Если значение пустое - пропускаем
            if not value_text:
                raw_fields.append({
                    "label": label_text,
                    "value": "",
                    "matched_key": keys[0] if keys else None
                })
                continue

            # Обработка составных полей (ИНН / КПП → два значения)
            if len(keys) > 1:
                # Пытаемся разделить значение
                values = re.split(r'[\s/,]+', value_text, maxsplit=len(keys)-1)
                values = [v.strip() for v in values if v.strip()]

                for i, key in enumerate(keys):
                    if key not in used_keys:
                        val = values[i] if i < len(values) else ""
                        if val:
                            extracted = self._extract_value(key, val)
                            normalized = self._normalize_value(key, extracted)
                            requisites[key] = normalized
                            used_keys.add(key)

                            # Валидация
                            warning = self._validate_value(key, normalized)
                            if warning:
                                warnings.append(warning)

                raw_fields.append({
                    "label": label_text,
                    "value": value_text,
                    "matched_key": "/".join(keys)
                })
            else:
                key = keys[0]
                if key not in used_keys:
                    extracted = self._extract_value(key, value_text)
                    normalized = self._normalize_value(key, extracted)
                    requisites[key] = normalized
                    used_keys.add(key)

                    # Валидация
                    warning = self._validate_value(key, normalized)
                    if warning:
                        warnings.append(warning)

                    raw_fields.append({
                        "label": label_text,
                        "value": value_text,
                        "matched_key": key
                    })

        # Fallback: парсинг свободного текста из всех ячеек таблицы
        # Собираем текст из всех ячеек, сохраняя порядок (сверху вниз, слева направо)
        all_cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and len(text) > 20:  # Только "длинные" ячейки со свободным текстом
                    all_cell_texts.append(text)

        # Парсим каждую ячейку отдельно, первые найденные значения приоритетнее
        for cell_text in all_cell_texts:
            free_text_result = self._extract_from_free_text(cell_text, used_keys)

            for key, value in free_text_result["requisites"].items():
                if key not in used_keys:
                    requisites[key] = value
                    used_keys.add(key)
                    raw_fields.append({
                        "label": f"[из текста]",
                        "value": value,
                        "matched_key": key
                    })

            warnings.extend(free_text_result["warnings"])

        return {
            "requisites": requisites,
            "raw_fields": raw_fields,
            "warnings": warnings,
            "table_found": True
        }

    def replace_placeholders(self, doc: Document, data: dict) -> int:
        """Замена плейсхолдеров {{key}} на значения"""
        count = 0

        for para in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))
                    count += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
                            count += 1

        return count

    def fill_table_by_labels(self, doc: Document, data: dict,
                             label_col: int = 1, value_col: int = 2,
                             skip_headers: bool = True) -> int:
        """Заполнение таблиц по меткам в столбце"""
        count = 0
        header_markers = ["№", "п/п", "наименование", "сведения"]

        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) > max(label_col, value_col):
                    if skip_headers and row_idx == 0:
                        continue

                    label_text = row.cells[label_col].text.strip().lower()

                    if skip_headers:
                        first_cell = row.cells[0].text.strip().lower()
                        if any(m in first_cell for m in header_markers):
                            continue

                    for label, json_key in self.label_mapping.items():
                        if label.lower() in label_text:
                            if json_key in data:
                                row.cells[value_col].text = str(data[json_key])
                                count += 1
                                break

        return count

    def fill(self, template_path: Path, output_path: Path, data: dict) -> dict:
        """
        Заполнение документа данными.
        Возвращает статистику заполнения:
        {
            "used_fields": {"field_name": count, ...},
            "unused_fields": ["field_name", ...]
        }
        """
        doc = Document(str(template_path))

        # Если postal_address не задан, используем address
        if "postal_address" not in data and "address" in data:
            data["postal_address"] = data["address"]

        # Собираем статистику использования полей
        used_fields: dict[str, int] = {}

        placeholder_used = self.replace_placeholders_with_stats(doc, data)
        label_used = self.fill_table_by_labels_with_stats(doc, data)

        # Объединяем статистику
        for key, count in placeholder_used.items():
            used_fields[key] = used_fields.get(key, 0) + count
        for key, count in label_used.items():
            used_fields[key] = used_fields.get(key, 0) + count

        doc.save(str(output_path))

        # Неиспользованные поля - те, что были в data но не использовались
        unused_fields = [k for k, v in data.items() if v and k not in used_fields]

        return {
            "used_fields": used_fields,
            "unused_fields": unused_fields
        }

    def replace_placeholders_with_stats(self, doc: Document, data: dict) -> dict[str, int]:
        """Замена плейсхолдеров {{key}} на значения, с подсчётом"""
        used: dict[str, int] = {}

        for para in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))
                    used[key] = used.get(key, 0) + 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
                            used[key] = used.get(key, 0) + 1

        return used

    def fill_table_by_labels_with_stats(self, doc: Document, data: dict,
                                        label_col: int = 1, value_col: int = 2,
                                        skip_headers: bool = True) -> dict[str, int]:
        """Заполнение таблиц по меткам в столбце, с подсчётом"""
        used: dict[str, int] = {}
        header_markers = ["№", "п/п", "наименование", "сведения"]

        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) > max(label_col, value_col):
                    if skip_headers and row_idx == 0:
                        continue

                    label_text = row.cells[label_col].text.strip().lower()

                    if skip_headers:
                        first_cell = row.cells[0].text.strip().lower()
                        if any(m in first_cell for m in header_markers):
                            continue

                    for label, json_key in self.label_mapping.items():
                        if label.lower() in label_text:
                            if json_key in data:
                                row.cells[value_col].text = str(data[json_key])
                                used[json_key] = used.get(json_key, 0) + 1
                                break

        return used


# Singleton instance
filler = DocxFiller()
