"""Конвертация docx в текстовое представление с сохранением структуры таблиц.

Используется для подготовки текста перед отправкой в LLM (v2).
"""
from pathlib import Path

from docx import Document

from app.services.docx_filler import DEFAULT_LABEL_MAPPING, COMPOSITE_FIELDS


def _table_to_fields_json(table) -> str:
    """Конвертация таблицы в структурированный JSON-список полей для LLM.

    Определяет колонку меток и колонку значений, пропускает заголовок.
    Возвращает JSON-массив объектов с row, col, field, current_value.
    """
    import json

    if not table.rows or len(table.rows) < 2:
        return "[]"

    num_cols = max(len(row.cells) for row in table.rows)
    if num_cols < 2:
        return "[]"

    # Колонка меток — обычно 1 (после № п/п), колонка значений — последняя
    label_col = 1 if num_cols >= 3 else 0
    value_col = num_cols - 1

    fields: list[dict] = []
    # Пропускаем строку 0 (заголовок)
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        label = row.cells[label_col].text.strip().replace("\n", " ") if label_col < len(row.cells) else ""
        current = row.cells[value_col].text.strip().replace("\n", " ") if value_col < len(row.cells) else ""

        if not label:
            continue

        fields.append({
            "row": row_idx,
            "col": value_col,
            "field": label,
            "current_value": current,
        })

    return json.dumps(fields, ensure_ascii=False, indent=2)


def _table_cells_matrix(table) -> list[list[str]]:
    """Извлечь матрицу текстов ячеек таблицы."""
    matrix: list[list[str]] = []
    for row in table.rows:
        row_cells = [cell.text.strip() for cell in row.cells]
        matrix.append(row_cells)
    return matrix


def docx_to_text(doc_path: Path) -> str:
    """Полный текст документа: параграфы + таблицы в Markdown-формате.

    Итерирует элементы документа в порядке появления (параграфы и таблицы),
    сохраняя порядок и переносы строк.
    """
    doc = Document(str(doc_path))

    # python-docx хранит body-элементы в doc.element.body;
    # чтобы сохранить порядок параграфов и таблиц, проходим по body children.
    from docx.oxml.ns import qn

    parts: list[str] = []

    # Построим индексы для быстрого маппинга XML-элемент -> объект
    para_map: dict[int, object] = {id(p._element): p for p in doc.paragraphs}
    table_map: dict[int, object] = {id(t._tbl): t for t in doc.tables}

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            para = para_map.get(id(child))
            if para is not None:
                text = para.text
                if text and text.strip():
                    parts.append(text)
        elif child.tag == qn("w:tbl"):
            table = table_map.get(id(child))
            if table is not None:
                md = _table_to_fields_json(table)
                if md:
                    parts.append(md)

    return "\n\n".join(parts)


def docx_tables_to_text(doc_path: Path) -> list[dict]:
    """Список таблиц с метаданными для обратного маппинга.

    Возвращает:
    [
        {
            "index": 0,
            "rows": 5,
            "cols": 3,
            "text": "| col1 | col2 | col3 |\\n| --- | --- | --- |\\n...",
            "cells": [["cell00", "cell01", "cell02"], ...]
        }
    ]
    """
    doc = Document(str(doc_path))
    result: list[dict] = []

    for idx, table in enumerate(doc.tables):
        cells = _table_cells_matrix(table)
        num_rows = len(table.rows)
        num_cols = max(len(row.cells) for row in table.rows) if table.rows else 0

        result.append({
            "index": idx,
            "rows": num_rows,
            "cols": num_cols,
            "text": _table_to_fields_json(table),
            "cells": cells,
        })

    return result


def find_requisites_table(doc_path: Path) -> dict | None:
    """Найти таблицу с наибольшим количеством заполняемых полей.

    Эвристика: 2-3 колонки, строки с текстом в первой колонке,
    максимальное число совпадений с паттернами из LABEL_MAPPING / COMPOSITE_FIELDS.

    Возвращает dict с метаданными таблицы (как в docx_tables_to_text) или None.
    """
    doc = Document(str(doc_path))

    if not doc.tables:
        return None

    all_patterns = list(COMPOSITE_FIELDS.keys()) + list(DEFAULT_LABEL_MAPPING.keys())

    best: dict | None = None
    best_score = 0

    for idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        num_cols = max(len(row.cells) for row in table.rows)

        # Эвристика: таблица реквизитов обычно имеет 2-3 колонки
        if num_cols < 2:
            continue

        score = 0
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip().lower()
                if not text:
                    continue
                for pattern in all_patterns:
                    if pattern in text:
                        score += 1
                        break  # Одно совпадение на ячейку

        if score > best_score:
            best_score = score
            cells = _table_cells_matrix(table)
            best = {
                "index": idx,
                "rows": len(table.rows),
                "cols": num_cols,
                "text": _table_to_fields_json(table),
                "cells": cells,
                "score": score,
            }

    # Возвращаем только если есть хотя бы несколько совпадений
    if best_score < 2:
        return None

    return best
