"""Тесты для fill endpoint (LLM-подход)"""
import io
import json
import pytest
from pathlib import Path
from unittest.mock import patch, AsyncMock, MagicMock

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


client = TestClient(app)


def _create_test_docx_with_table(path: Path, rows: int = 3, cols: int = 2) -> None:
    """Создать тестовый docx с таблицей"""
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    # Заполняем первую колонку метками
    labels = ["ИНН", "КПП", "Наименование организации"]
    for i, label in enumerate(labels[:rows]):
        table.rows[i].cells[0].text = label
        # Вторая колонка пустая (для заполнения)
        table.rows[i].cells[1].text = ""
    doc.save(str(path))


def _create_test_docx_no_table(path: Path) -> None:
    """Создать тестовый docx без таблиц"""
    doc = Document()
    doc.add_paragraph("Просто текст без таблиц")
    doc.save(str(path))


def _docx_to_bytes(path: Path) -> bytes:
    """Прочитать docx в bytes"""
    with open(path, "rb") as f:
        return f.read()


@pytest.fixture
def tmp_docx_with_table(tmp_path):
    """Создать временный docx с таблицей"""
    docx_path = tmp_path / "test_template.docx"
    _create_test_docx_with_table(docx_path)
    return docx_path


@pytest.fixture
def tmp_docx_no_table(tmp_path):
    """Создать временный docx без таблиц"""
    docx_path = tmp_path / "test_no_table.docx"
    _create_test_docx_no_table(docx_path)
    return docx_path


class TestFillSuccess:
    """Тест успешного заполнения через LLM"""

    def test_fill_success(self, tmp_docx_with_table):
        """Мок LLM возвращает инструкции -> документ заполняется"""
        docx_bytes = _docx_to_bytes(tmp_docx_with_table)
        requisites_json = json.dumps({"inn": "1234567890", "kpp": "123456789"})

        mock_instructions = [
            {"row": 0, "col": 1, "value": "1234567890"},
            {"row": 1, "col": 1, "value": "123456789"},
        ]

        with patch(
            "app.api.routes.find_requisites_table",
            return_value={
                "index": 0,
                "rows": 3,
                "cols": 2,
                "text": "| ИНН | |\n| --- | --- |\n| КПП | |",
                "cells": [["ИНН", ""], ["КПП", ""], ["Наименование организации", ""]],
            },
        ), patch(
            "app.api.routes.llm_service"
        ) as mock_llm:
            mock_llm.generate_fill_instructions = AsyncMock(
                return_value=mock_instructions
            )

            response = client.post(
                "/api/fill",
                files={"file": ("template.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"requisites": requisites_json},
            )

        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["filled_fields"] == 2
        assert data["total_instructions"] == 2
        assert "download_url" in data
        assert data["filename"].endswith(".docx")
        assert "Заполнено 2 ячеек" in data["message"]


class TestFillNoTable:
    """Тест: документ без таблиц"""

    def test_fill_no_table(self, tmp_docx_no_table):
        """Документ без таблиц -> HTTP 400"""
        docx_bytes = _docx_to_bytes(tmp_docx_no_table)
        requisites_json = json.dumps({"inn": "1234567890"})

        with patch(
            "app.api.routes.find_requisites_table",
            return_value=None,
        ), patch(
            "app.api.routes.docx_tables_to_text",
            return_value=[],
        ):
            response = client.post(
                "/api/fill",
                files={"file": ("template.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"requisites": requisites_json},
            )

        assert response.status_code == 400
        assert "Таблица для заполнения не найдена" in response.json()["detail"]


class TestFillLLMError:
    """Тест: ошибка LLM -> HTTP 502"""

    def test_fill_llm_timeout(self, tmp_docx_with_table):
        """Таймаут LLM API -> HTTP 502"""
        import openai

        docx_bytes = _docx_to_bytes(tmp_docx_with_table)
        requisites_json = json.dumps({"inn": "1234567890"})

        with patch(
            "app.api.routes.find_requisites_table",
            return_value={
                "index": 0,
                "rows": 3,
                "cols": 2,
                "text": "| ИНН | |",
                "cells": [["ИНН", ""]],
            },
        ), patch(
            "app.api.routes.llm_service"
        ) as mock_llm:
            mock_llm.generate_fill_instructions = AsyncMock(
                side_effect=openai.APITimeoutError(request=MagicMock())
            )

            response = client.post(
                "/api/fill",
                files={"file": ("template.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"requisites": requisites_json},
            )

        assert response.status_code == 502
        assert "Таймаут LLM API" in response.json()["detail"]

    def test_fill_llm_value_error(self, tmp_docx_with_table):
        """Невалидный ответ LLM -> HTTP 502"""
        docx_bytes = _docx_to_bytes(tmp_docx_with_table)
        requisites_json = json.dumps({"inn": "1234567890"})

        with patch(
            "app.api.routes.find_requisites_table",
            return_value={
                "index": 0,
                "rows": 3,
                "cols": 2,
                "text": "| ИНН | |",
                "cells": [["ИНН", ""]],
            },
        ), patch(
            "app.api.routes.llm_service"
        ) as mock_llm:
            mock_llm.generate_fill_instructions = AsyncMock(
                side_effect=ValueError("JSON-массив не найден в ответе LLM")
            )

            response = client.post(
                "/api/fill",
                files={"file": ("template.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"requisites": requisites_json},
            )

        assert response.status_code == 502
        assert "Ошибка разбора ответа LLM" in response.json()["detail"]


class TestFillInvalidInstructions:
    """Тест: инструкции с row/col за пределами таблицы -> пропускаются"""

    def test_fill_invalid_instructions_skipped(self, tmp_docx_with_table):
        """Инструкции за пределами таблицы пропускаются, валидные применяются"""
        docx_bytes = _docx_to_bytes(tmp_docx_with_table)
        requisites_json = json.dumps({"inn": "1234567890"})

        mock_instructions = [
            {"row": 0, "col": 1, "value": "1234567890"},  # валидная
            {"row": 99, "col": 1, "value": "out_of_range"},  # за пределами
            {"row": 1, "col": 99, "value": "out_of_range"},  # за пределами
        ]

        with patch(
            "app.api.routes.find_requisites_table",
            return_value={
                "index": 0,
                "rows": 3,
                "cols": 2,
                "text": "| ИНН | |\n| --- | --- |",
                "cells": [["ИНН", ""], ["КПП", ""], ["Наименование", ""]],
            },
        ), patch(
            "app.api.routes.llm_service"
        ) as mock_llm:
            mock_llm.generate_fill_instructions = AsyncMock(
                return_value=mock_instructions
            )

            response = client.post(
                "/api/fill",
                files={"file": ("template.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"requisites": requisites_json},
            )

        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        # Только 1 из 3 инструкций валидна
        assert data["filled_fields"] == 1
        # total_instructions = кол-во строк с метками в таблице (row 0 — заголовок, пропускается)
        assert data["total_instructions"] == 2
        assert data["skipped_count"] == 2


class TestRequisitesToXml:
    """Тесты для утилиты requisites_to_xml"""

    def test_basic(self):
        from app.api.routes import requisites_to_xml

        result = requisites_to_xml({"inn": "123", "kpp": "456"})
        assert "<requisites>" in result
        assert '</requisites>' in result
        assert '<field name="inn">123</field>' in result
        assert '<field name="kpp">456</field>' in result

    def test_empty(self):
        from app.api.routes import requisites_to_xml

        result = requisites_to_xml({})
        assert result == "<requisites>\n</requisites>"
