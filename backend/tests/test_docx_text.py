"""Тесты для backend/app/services/docx_text.py"""
from pathlib import Path

import pytest
from docx import Document

from app.services.docx_text import docx_to_text, docx_tables_to_text, find_requisites_table


def _make_docx(tmp_path: Path, name: str = "test.docx") -> Path:
    """Создать пустой docx и вернуть путь."""
    path = tmp_path / name
    doc = Document()
    doc.save(str(path))
    return path


# ---------- docx_to_text ----------


def test_docx_to_text_paragraphs(tmp_path: Path):
    """Параграфы конвертируются в plain text с переносами."""
    path = tmp_path / "para.docx"
    doc = Document()
    doc.add_paragraph("Первая строка")
    doc.add_paragraph("Вторая строка")
    doc.add_paragraph("Третья строка")
    doc.save(str(path))

    text = docx_to_text(path)
    assert "Первая строка" in text
    assert "Вторая строка" in text
    assert "Третья строка" in text


def test_docx_to_text_with_table(tmp_path: Path):
    """Таблица конвертируется в Markdown pipe-формат."""
    path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Заголовок документа")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(0, 2).text = "C1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    table.cell(1, 2).text = "C2"
    doc.save(str(path))

    text = docx_to_text(path)
    assert "Заголовок документа" in text
    assert "| A1 | B1 | C1 |" in text
    assert "| --- | --- | --- |" in text
    assert "| A2 | B2 | C2 |" in text


def test_docx_to_text_preserves_order(tmp_path: Path):
    """Порядок параграфов и таблиц сохраняется."""
    path = tmp_path / "order.docx"
    doc = Document()
    doc.add_paragraph("ДО ТАБЛИЦЫ")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "ЯЧЕЙКА"
    table.cell(0, 1).text = "ЗНАЧЕНИЕ"
    doc.add_paragraph("ПОСЛЕ ТАБЛИЦЫ")
    doc.save(str(path))

    text = docx_to_text(path)
    pos_before = text.index("ДО ТАБЛИЦЫ")
    pos_cell = text.index("ЯЧЕЙКА")
    pos_after = text.index("ПОСЛЕ ТАБЛИЦЫ")
    assert pos_before < pos_cell < pos_after


def test_empty_document(tmp_path: Path):
    """Пустой docx возвращает пустую строку."""
    path = _make_docx(tmp_path)
    text = docx_to_text(path)
    assert text == ""


# ---------- docx_tables_to_text ----------


def test_docx_tables_to_text_metadata(tmp_path: Path):
    """Метаданные таблицы: index, rows, cols, text, cells."""
    path = tmp_path / "meta.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Метка"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "ИНН"
    table.cell(1, 1).text = "1234567890"
    table.cell(2, 0).text = "КПП"
    table.cell(2, 1).text = "123456789"
    doc.save(str(path))

    tables = docx_tables_to_text(path)
    assert len(tables) == 1

    t = tables[0]
    assert t["index"] == 0
    assert t["rows"] == 3
    assert t["cols"] == 2
    assert "| Метка | Значение |" in t["text"]
    assert "| --- | --- |" in t["text"]
    assert t["cells"][1] == ["ИНН", "1234567890"]
    assert t["cells"][2] == ["КПП", "123456789"]


def test_docx_tables_to_text_multiple_tables(tmp_path: Path):
    """Несколько таблиц возвращают правильные индексы."""
    path = tmp_path / "multi.docx"
    doc = Document()
    doc.add_table(rows=1, cols=2)
    doc.add_table(rows=2, cols=3)
    doc.save(str(path))

    tables = docx_tables_to_text(path)
    assert len(tables) == 2
    assert tables[0]["index"] == 0
    assert tables[0]["cols"] == 2
    assert tables[1]["index"] == 1
    assert tables[1]["cols"] == 3


def test_docx_tables_to_text_no_tables(tmp_path: Path):
    """Документ без таблиц возвращает пустой список."""
    path = _make_docx(tmp_path)
    tables = docx_tables_to_text(path)
    assert tables == []


# ---------- find_requisites_table ----------


def test_find_requisites_table_correct(tmp_path: Path):
    """Находит таблицу с реквизитами среди нескольких."""
    path = tmp_path / "req.docx"
    doc = Document()

    # Таблица 1 — нерелевантная
    t1 = doc.add_table(rows=2, cols=2)
    t1.cell(0, 0).text = "Товар"
    t1.cell(0, 1).text = "Цена"
    t1.cell(1, 0).text = "Молоко"
    t1.cell(1, 1).text = "100"

    # Таблица 2 — реквизиты
    t2 = doc.add_table(rows=5, cols=2)
    t2.cell(0, 0).text = "Наименование организации"
    t2.cell(0, 1).text = ""
    t2.cell(1, 0).text = "ИНН"
    t2.cell(1, 1).text = ""
    t2.cell(2, 0).text = "КПП"
    t2.cell(2, 1).text = ""
    t2.cell(3, 0).text = "Расчетный счет"
    t2.cell(3, 1).text = ""
    t2.cell(4, 0).text = "БИК"
    t2.cell(4, 1).text = ""

    doc.save(str(path))

    result = find_requisites_table(path)
    assert result is not None
    assert result["index"] == 1
    assert result["rows"] == 5
    assert result["cols"] == 2
    assert result["score"] >= 4


def test_find_requisites_table_none(tmp_path: Path):
    """Возвращает None если нет таблиц с реквизитами."""
    path = tmp_path / "noreq.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Товар"
    t.cell(0, 1).text = "Цена"
    t.cell(1, 0).text = "Молоко"
    t.cell(1, 1).text = "100"
    doc.save(str(path))

    result = find_requisites_table(path)
    assert result is None


def test_find_requisites_table_no_tables(tmp_path: Path):
    """Документ без таблиц возвращает None."""
    path = _make_docx(tmp_path)
    result = find_requisites_table(path)
    assert result is None


def test_find_requisites_table_three_columns(tmp_path: Path):
    """Таблица с 3 колонками (номер, метка, значение) распознаётся."""
    path = tmp_path / "threecol.docx"
    doc = Document()
    t = doc.add_table(rows=4, cols=3)
    t.cell(0, 0).text = "1"
    t.cell(0, 1).text = "Полное наименование"
    t.cell(0, 2).text = ""
    t.cell(1, 0).text = "2"
    t.cell(1, 1).text = "ИНН"
    t.cell(1, 2).text = ""
    t.cell(2, 0).text = "3"
    t.cell(2, 1).text = "ОГРН"
    t.cell(2, 2).text = ""
    t.cell(3, 0).text = "4"
    t.cell(3, 1).text = "Телефон"
    t.cell(3, 2).text = ""
    doc.save(str(path))

    result = find_requisites_table(path)
    assert result is not None
    assert result["cols"] == 3
    assert result["score"] >= 3
